"""Deterministic local text and tabular metadata extraction."""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from utils import ROOT, assert_raw_read_only


def extract_plain(path: Path) -> dict[str, Any]:
    return {"kind": "text", "text": path.read_text(encoding="utf-8", errors="replace")}


def extract_pdf(path: Path) -> dict[str, Any]:
    try:
        import fitz
    except ImportError:
        pages = extract_pdf_with_pdfkit(path)
    else:
        pages = []
        with fitz.open(path) as document:
            for number, page in enumerate(document, start=1):
                pages.append({"page": number, "text": page.get_text("text")})
    return {
        "kind": "pdf",
        "page_count": len(pages),
        "pages": pages,
        "text": "\n\n".join(
            f"--- Page {item['page']} ---\n{item['text']}" for item in pages
        ),
    }


def extract_pdf_with_pdfkit(path: Path) -> list[dict[str, Any]]:
    """Use macOS PDFKit when PyMuPDF is not installed."""
    swift_source = ROOT / "scripts" / "extract_pdf_pdfkit.swift"
    if not swift_source.exists():
        raise RuntimeError("PDF extraction requires pymupdf or the PDFKit helper")
    binary = Path(tempfile.gettempdir()) / "researchai_pdfkit_extract"
    if not binary.exists() or binary.stat().st_mtime < swift_source.stat().st_mtime:
        env = os.environ.copy()
        env["SWIFT_MODULECACHE_PATH"] = str(Path(tempfile.gettempdir()) / "swift-module-cache")
        env["CLANG_MODULE_CACHE_PATH"] = str(Path(tempfile.gettempdir()) / "clang-module-cache")
        subprocess.run(
            ["swiftc", str(swift_source), "-o", str(binary)],
            check=True,
            env=env,
            capture_output=True,
            text=True,
        )
    result = subprocess.run(
        [str(binary), str(path)],
        check=True,
        capture_output=True,
        text=True,
    )
    parts = re.split(r"<<<PDFKIT_PAGE:(\d+)>>>\n", result.stdout)[1:]
    return [
        {"page": int(parts[index]), "text": parts[index + 1]}
        for index in range(0, len(parts), 2)
    ]


def extract_docx(path: Path) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX extraction requires python-docx") from exc
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    tables = []
    for table in document.tables:
        tables.append([[cell.text for cell in row.cells] for row in table.rows])
    return {
        "kind": "docx",
        "paragraph_count": len(paragraphs),
        "tables": tables,
        "text": "\n".join(paragraphs),
    }


def dataframe_summary(frame: Any, source_type: str) -> dict[str, Any]:
    numeric = frame.describe(include="number").to_dict() if not frame.empty else {}
    categorical = {
        str(column): int(frame[column].nunique(dropna=True))
        for column in frame.select_dtypes(exclude="number").columns
    }
    target_tokens = ("target", "label", "class", "outcome", "reward", "delivery", "delay")
    target_candidates = [
        str(column)
        for column in frame.columns
        if any(token in str(column).casefold() for token in target_tokens)
    ]
    summary = {
        "kind": source_type,
        "row_count": int(frame.shape[0]),
        "column_count": int(frame.shape[1]),
        "column_names": [str(column) for column in frame.columns],
        "dtypes": {str(key): str(value) for key, value in frame.dtypes.items()},
        "missing_values": {
            str(key): int(value) for key, value in frame.isna().sum().items()
        },
        "numeric_describe": numeric,
        "categorical_cardinality": categorical,
        "target_column_candidates": target_candidates,
        "research_relevance_note": (
            "TODO: assess columns, units, collection protocol, leakage risk, "
            "and relevance to FANET routing experiments."
        ),
    }
    summary["text"] = json.dumps(summary, ensure_ascii=False, indent=2, default=str)
    return summary


def extract_table(path: Path) -> dict[str, Any]:
    try:
        import pandas as pd
    except ImportError as exc:
        raise RuntimeError("Tabular extraction requires pandas") from exc
    if path.suffix.casefold() == ".csv":
        frame = pd.read_csv(path)
        return dataframe_summary(frame, "csv")
    frame = pd.read_excel(path)
    return dataframe_summary(frame, "xlsx")


def extract(path: Path, require_raw: bool = True) -> dict[str, Any]:
    path = path.resolve()
    if require_raw:
        assert_raw_read_only(path)
    suffix = path.suffix.casefold()
    if suffix in {".txt", ".md"}:
        return extract_plain(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".csv", ".xlsx"}:
        return extract_table(path)
    raise ValueError(f"Unsupported file type: {suffix or '<none>'}")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("path", type=Path)
    parser.add_argument("--allow-non-raw", action="store_true")
    args = parser.parse_args()
    result = extract(args.path, require_raw=not args.allow_non_raw)
    print(json.dumps(result, ensure_ascii=False, indent=2, default=str))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
